import os
import csv

os.environ["ANONYMIZED_TELEMETRY"] = "false"
os.environ["DISABLE_TELEMETRY"] = "1"
os.environ["CHROMA_TELEMETRY_ENABLED"] = "false"

import json
import logging
import threading
import pandas as pd
from typing import Any, Dict, List, Optional

# langchain
from langchain.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    SystemMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain_core.messages import (
    HumanMessage,
    AIMessage,
)
from langchain_ollama import OllamaLLM, OllamaEmbeddings

# llama-index & chroma
import chromadb
from llama_index.core import Settings  # 全局
from llama_index.core import Document
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, StorageContext
from llama_index.vector_stores.chroma import ChromaVectorStore

# 日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import re
from docx import Document as DocxDocument
from PyPDF2 import PdfReader


class TopKLogSystem:
    """
    基于 DeepSeek-R1:7B 的日志分析系统

    使用模型:
    - LLM: DeepSeek-R1:7B (deepseek-r1:7b)
      * 架构: 基于 Qwen2 架构的 DeepSeek-R1 模型
      * 参数量: 7.6B
      * 上下文长度: 131072 tokens
      * 特性: 支持思考过程 (thinking)，使用 <think> 标签
      * Temperature: 0.1 (低温度保证输出稳定性)
    - Embedding: BGE-Large (bge-large:latest)
      * 用于向量检索和文档嵌入
    """

    def __init__(
        self,
        log_path: str,
        llm: str,
        embedding_model: str,
    ) -> None:
        # 初始化嵌入模型 (BGE-Large)
        self.embedding_model = OllamaEmbeddings(model=embedding_model)
        self._default_llm_name = llm
        self._llm_cache: Dict[str, OllamaLLM] = {}
        self._llm_lock = threading.RLock()
        # 初始化大语言模型 (DeepSeek-R1:7B)
        self.llm = self._get_or_create_llm(llm)

        Settings.llm = self.llm
        Settings.embed_model = self.embedding_model

        self.log_path = log_path
        self.log_index = None
        self.vector_store = None
        self._build_vectorstore()

    # 构建向量数据库的核心函数
    def _build_vectorstore(self):
        vector_store_path = "./data/vector_stores"

        # 检查 vector_stores 文件夹是否存在
        if os.path.exists(vector_store_path):
            logger.info(f"向量数据库文件夹已存在，加载现有索引: {vector_store_path}")

            try:
                # 1. 连接到现有的 ChromaDB
                chroma_client = chromadb.PersistentClient(path=vector_store_path)

                # 2. 获取集合
                log_collection = chroma_client.get_collection("log_collection")

                # 3. 实例化 LlamaIndex 的 VectorStore
                log_vector_store = ChromaVectorStore(chroma_collection=log_collection)

                # 4.从 VectorStore 加载索引
                self.log_index = VectorStoreIndex.from_vector_store(
                    vector_store=log_vector_store
                )
                self.vector_store = log_vector_store

                logger.info("成功从现有数据库加载索引。")

            except Exception as e:
                logger.error(f"加载现有向量数据库失败: {e}. 系统将无法进行日志检索。")

            return  # 结束函数

        logger.info(f"向量数据库文件夹不存在，开始构建: {vector_store_path}")
        os.makedirs(vector_store_path, exist_ok=True)

        chroma_client = chromadb.PersistentClient(path=vector_store_path)
        log_collection = chroma_client.get_or_create_collection("log_collection")

        log_vector_store = ChromaVectorStore(chroma_collection=log_collection)
        self.vector_store = log_vector_store  # 保持一致性

        log_storage_context = StorageContext.from_defaults(
            vector_store=log_vector_store
        )
        if log_documents := self._load_documents(self.log_path):
            self.log_index = VectorStoreIndex.from_documents(
                log_documents,
                storage_context=log_storage_context,
                show_progress=True,
            )
            logger.info(f"日志库索引构建完成，共 {len(log_documents)} 条日志")
        else:
            logger.info("未加载到任何日志文档，向量数据库未更新")

    # 函数用来读取文档,添加可读取文档类型,并支持遍历子文件夹下的文件
    @staticmethod
    def _load_documents(data_path: str) -> List[Document]:
        """
        递归遍历 data_path 下所有文件（包括子文件夹），加载支持的文档类型。
        """
        if not os.path.exists(data_path):
            logger.warning(f"数据路径不存在: {data_path}")
            return []
        documents = []
        # 使用 os.walk 递归遍历所有文件
        for root, dirs, files in os.walk(data_path):
            for file in files:
                ext = os.path.splitext(file)[1]
                if ext not in [
                    ".txt",
                    ".md",
                    ".json",
                    ".jsonl",
                    ".csv",
                    ".log",
                    ".xml",
                    ".yaml",
                    ".yml",
                    ".docx",
                    ".pdf",
                ]:
                    continue
                file_path = os.path.join(root, file)
                try:
                    if ext == ".csv":
                        documents.extend(TopKLogSystem._process_csv(file_path))
                    elif ext in [".json", ".jsonl"]:
                        documents.extend(TopKLogSystem._process_json(file_path, ext))
                    elif ext in [".yaml", ".yml"]:
                        documents.extend(TopKLogSystem._process_yaml(file_path))
                    elif ext == ".xml":
                        documents.extend(TopKLogSystem._process_xml(file_path))
                    elif ext == ".log":
                        documents.extend(TopKLogSystem._process_log(file_path))
                    elif ext == ".docx":
                        documents.extend(TopKLogSystem._process_docx(file_path))
                    elif ext == ".pdf":
                        documents.extend(TopKLogSystem._process_pdf(file_path))
                    else:
                        documents.extend(TopKLogSystem._process_text(file_path))
                except Exception as e:
                    logger.error(f"加载文档失败 {file_path}: {e}")
        return documents

    # 各种文件类型的处理函数
    @staticmethod
    def _process_csv(file_path: str) -> List[Document]:
        documents = []
        chunk_size = 1000
        for chunk in pd.read_csv(file_path, chunksize=chunk_size, on_bad_lines="skip"):
            for row in chunk.itertuples(index=False):
                content = str(row).replace("Pandas", " ")
                documents.append(Document(text=content))
        return documents

    @staticmethod
    def _process_json(file_path: str, ext: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            if ext == ".json":
                data = json.load(f)
                documents.append(Document(text=json.dumps(data, ensure_ascii=False)))
            elif ext == ".jsonl":
                for line in f:
                    data = json.loads(line.strip())
                    documents.append(
                        Document(text=json.dumps(data, ensure_ascii=False))
                    )
        return documents

    @staticmethod
    def _process_yaml(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            documents.append(Document(text=content))
        return documents

    @staticmethod
    def _process_xml(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            documents.append(Document(text=content))
        return documents

    @staticmethod
    def _process_log(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line in f:
                documents.append(Document(text=line.strip()))
        return documents

    @staticmethod
    def _process_docx(file_path: str) -> List[Document]:
        documents = []
        doc = DocxDocument(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                documents.append(Document(text=paragraph.text.strip()))
        return documents

    @staticmethod
    def _process_pdf(file_path: str) -> List[Document]:
        documents = []
        reader = PdfReader(file_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                documents.append(Document(text=text.strip()))
        return documents

    @staticmethod
    def _process_text(file_path: str) -> List[Document]:
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            documents.append(Document(text=content))
        return documents

    def retrieve_logs(self, query: str, top_k: int = 10) -> List[Dict]:
        if not self.log_index:
            logger.warning("Log index 未初始化，跳过检索。")
            return []
        try:
            retriever = self.log_index.as_retriever(similarity_top_k=top_k)
            results = retriever.retrieve(query)
            formatted_results = []
            for result in results:
                formatted_results.append(
                    {"content": result.text, "score": result.score}
                )
            return formatted_results
        except Exception as e:
            logger.error(f"日志检索失败: {e}")
            return []

    # (修改) context 现在是一个字典
    def _get_or_create_llm(self, model_name: Optional[str]) -> OllamaLLM:
        target_name = (model_name or self._default_llm_name or "").strip()
        if not target_name:
            target_name = self._default_llm_name

        cached_llm = self._llm_cache.get(target_name)
        if cached_llm is not None:
            return cached_llm

        new_llm = OllamaLLM(model=target_name, temperature=0.1)
        self._llm_cache[target_name] = new_llm
        return new_llm

    def generate_response(
        self,
        query: str,
        context: Dict,
        history: List[Dict] = None,
        model_name: Optional[str] = None,
    ):
        prompt_messages = self._build_prompt(query, context, history)

        llm_to_use = self._get_or_create_llm(model_name)

        # 更新当前使用的 LLM，确保后续依赖 Settings.llm 的流程保持一致
        with self._llm_lock:
            self.llm = llm_to_use
            Settings.llm = llm_to_use

        try:
            for chunk in llm_to_use.stream(prompt_messages):
                yield chunk

        except Exception as e:
            logger.error(f"LLM调用失败: {e}")
            yield f"生成响应时出错: {str(e)}"

    # (修改) context 现在是一个字典, 并更新 Prompt
    def _build_prompt(
        self, query: str, context: Dict, history: List[Dict] = None
    ) -> List:

        system_message = SystemMessagePromptTemplate.from_template(
            """
            你是一个多任务SRE助手。你的首要任务是 **[判断意图]**，然后根据意图选择正确的 **[响应模式]**。

            你有两种响应模式：
            1.  **[SRE分析模式]**: 当用户的问题与故障排查、日志分析、系统错误相关时使用。
            2.  **[常规对话模式]**: 当用户进行常规闲聊 (如 "你好")、历史回顾 (如 "我刚才问了什么") 或提出与日志无关的问题 (如 "介绍一下天津大学") 时使用。
            
            **[!!! 可用工具 (SRE模式专用) !!!]**
            你现在有两种工具上下文：
            1.  **[日志数据库 (Log DB)]**: 包含本地的、详细的系统日志。
            2.  **[联网搜索 (Web Search)]**: 包含来自互联网的实时信息。

            你的回答必须遵循以下质量要求：
            1.  **专业严谨**：在 [SRE分析模式] 下，你的分析必须基于上下文（[日志数据库] 和/或 [联网搜索]），严禁凭空猜测。
            2.  **优先使用日志**：如果 [日志数据库] 提供了足够的信息，优先使用它。只有当日志信息不足或用户明确询问需要外部知识时，才使用 [联网搜索]。
            3.  **清晰可读**：使用 Markdown 格式（如列表、代码块、粗体）来组织你的回答。
            4.  **上下文感知**：你必须能够 **自主判断** 是否需要结合 **历史对话** 来理解用户的真实意图或SRE问题。

            **[!!! 绝对指令：输出格式 !!!]**
            1.  你 **必须** 且 **只能** 使用 `<think>...</think>` 标签来包裹你的所有内部思考步骤 (包括意图分析、SRE分析框架等)。
            2.  在 `<think>...</think>` 标签之外，你 **必须** 且 **只能** 输出 **最终的、直接面向用户** 的回复。
            3.  最终回复中 **严禁** 包含 "步骤 1"、"步骤 2"、"意图分析"、"根本原因"、"最终回复草稿" 等任何思考过程的字样。
            """
        )

        # (修改) 2. 准备日志上下文 (Log DB)
        log_context_str = "## [可用工具 1: 日志数据库 (Log DB)]\n"
        log_data = context.get("log_context", [])  # 从字典获取
        if not log_data:
            log_context_str += "（未从日志数据库检索到相关内容）\n"
        else:
            for i, log in enumerate(log_data, 1):
                # 确保 score 是浮点数以便格式化
                score = log.get("score", 0.0)
                log_context_str += f"日志 {i} (Score: {score:.2f}): {log['content']}\n"

        # (新增) 3. 准备联网搜索上下文 (Web Search)
        web_context_str = "## [可用工具 2: 联网搜索 (Web Search)]\n"
        web_data = context.get("web_context", [])
        if not web_data:
            web_context_str += "（未启用或未从联网搜索检索到相关内容）\n"
        else:
            for i, web_result in enumerate(web_data, 1):
                web_context_str += f"网页 {i} (Source: {web_result.get('source', 'N/A')}): {web_result['content']}\n"

        # (修改) 4. 更新用户模板
        user_message_template = HumanMessagePromptTemplate.from_template(
            """
            {log_context}
            {web_context}
            
            ---
            ## [任务] 当前用户问题:
            {query}
            ---

            ## [执行指令]
            1.  在 <think> 块中严格按照SRE分析框架进行思考。
            2.  在 <think> 块 **之外**，生成最终的、面向用户的回复。

            <think>
            
            **步骤 1: 意图分析 (Intent Analysis)**
            * 用户当前问题是："{query}"
            * 历史对话上下文分析：(简要评估历史对话。例如：用户是在追问SRE问题、转换话题、提供新信息，还是在进行常规对话？)
            * 意图判断：(基于 **当前问题** 和 **历史上下文**，填写 [SRE分析模式] 或 [常规对话模式])

            **步骤 2: 响应策略 (Response Strategy)**
            * 如果 (If) 意图是 [常规对话模式]:
                * 我将生成一个友好、对应的回复。我将忽略 [可用工具] 中的所有上下文。
            * 如果 (If) 意图是 [SRE分析模式]:
                * 我必须使用下面的 [SRE分析框架] 来分析 [可用工具] 中的 [日志数据库] 和/或 [联网搜索]，并结合 **历史对话** 和 **用户问题** 进行回答。

            **步骤 3: SRE分析框架 (仅SRE模式执行)**
            * **a. 问题现象 (Symptom):**
                * (结合 {query} **以及历史对话**，总结用户描述的完整问题现象。)
            * **b. 上下文关联 (Context Correlation):**
                * (审查 [可用工具 1: 日志数据库]。哪些日志与 [a. 问题现象] 相关？)
                * (审查 [可用工具 2: 联网搜索]。哪些网页结果与 [a. 问题现象] 相关？)
                * (判断：日志信息是否足够？是否需要 [联网搜索] 来补充背景知识或解决方案？优先使用日志。)
            * **c. 根本原因 (Root Cause Hypothesis):**
                * (基于 [b. 上下文关联]、{query} **和历史对话**，提出1-2个最可能的根本原因。优先使用日志，辅以网页搜索。)
            * **d. 建议方案 (Actionable Steps):**
                * (提出具体的解决步骤或进一步的排查指令。)

            **步骤 4: 回复构思 (Response Rationale)**
            * (基于 [步骤2] 或 [步骤3] 的分析，在此处 **构思** 给用户的最终回复，使用详细且严谨的专家风格。这 **不是** 最终草稿。)
            
            </think>
            
            """
        )

        prompt_template = ChatPromptTemplate.from_messages(
            [
                system_message,
                MessagesPlaceholder(variable_name="chat_history"),
                user_message_template,
            ]
        )

        formatted_history = []
        if history:
            for msg in history:
                if msg["role"] == "user":
                    formatted_history.append(HumanMessage(content=msg["content"]))
                elif msg["role"] == "assistant":
                    clean_content = re.sub(
                        r"<think>.*?</think>\s*",
                        "",
                        msg["content"],
                        flags=re.DOTALL,
                    )
                    formatted_history.append(AIMessage(content=clean_content.strip()))
                else:
                    # 兼容旧格式 (如果存在)
                    formatted_history.append(AIMessage(content=msg["content"]))

        # (修改) 传递 web_context
        return prompt_template.format_prompt(
            chat_history=formatted_history,
            log_context=log_context_str,
            web_context=web_context_str,
            query=query,
        ).to_messages()

    # (修改) 更新 query 方法以适应新的 generate_response 签名 (主要用于内部测试)
    def query(
        self,
        query: str,
        history: List[Dict] = None,
        use_db_search: bool = True,
        use_web_search: bool = False,
        model_name: Optional[str] = None,
    ):
        log_results = []
        if use_db_search:
            log_results = self.retrieve_logs(query)

        web_results = []
        if use_web_search:
            # 内部测试无法调用 services.py 的 mock，这里简单模拟
            logger.info(f"[MOCK-QUERY] 联网搜索: {query}")
            web_results = [{"content": "模拟网页结果", "source": "mock.com"}]

        combined_context = {"log_context": log_results, "web_context": web_results}

        for chunk in self.generate_response(
            query,
            combined_context,
            history,
            model_name=model_name,
        ):
            yield chunk


if __name__ == "__main__":
    # 测试使用 DeepSeek-R1:7B 和 BGE-Large 嵌入模型
    system = TopKLogSystem(
        log_path="./data/log",
        llm="deepseek-r1:7b",  # DeepSeek-R1:7B - 基于 Qwen2 架构
        embedding_model="bge-large:latest",  # BGE-Large 嵌入模型
    )

    query1 = "我遇到了数据库问题"
    print("查询1:", query1)
    print("响应1 (流式 - 仅数据库):")
    full_response_1 = ""
    # (修改) 测试调用
    for chunk in system.query(query1, use_db_search=True, use_web_search=False):
        print(chunk, end="", flush=True)
        full_response_1 += chunk
    print("\n--- 流结束 ---")

    history_example = [
        {"role": "user", "content": query1},
        {"role": "assistant", "content": full_response_1},
    ]

    query2 = "我刚才问了什么？"
    print("\n查询2 (测试历史对话):", query2)
    print("响应2 (流式 - 无搜索):")
    full_response_2 = ""
    for chunk in system.query(
        query2, history=history_example, use_db_search=False, use_web_search=False
    ):
        print(chunk, end="", flush=True)
        full_response_2 += chunk
    print("\n--- 流结束 ---")

    query3 = "是连接池耗尽的问题，如何解决？"
    history_example.append({"role": "user", "content": query2})
    history_example.append({"role": "assistant", "content": full_response_2})
    print("\n查询3 (测试日志分析 + 联网):", query3)
    print("响应3 (流式):")
    full_response_3 = ""
    for chunk in system.query(
        query3, history=history_example, use_db_search=True, use_web_search=True
    ):
        print(chunk, end="", flush=True)
        full_response_3 += chunk
    print("\n--- 流结束 ---")

    query4 = "你好"
    history_example.append({"role": "user", "content": query3})
    history_example.append({"role": "assistant", "content": full_response_3})
    print("\n查询4 (测试常规对话):", query4)
    print("响应4 (流式):")
    for chunk in system.query(
        query4, history=history_example, use_db_search=False, use_web_search=False
    ):
        print(chunk, end="", flush=True)
    print("\n--- 流结束 ---")
